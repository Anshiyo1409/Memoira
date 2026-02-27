# doc_utils.py
from PyPDF2 import PdfReader

def extract_text_from_file(file):
    """
    Extract text from PDF, TXT, or DOCX.
    Accepts Streamlit UploadedFile or file path string.
    """
    if hasattr(file, "name"):
        ext = file.name.split(".")[-1].lower()
    else:
        ext = file.split(".")[-1].lower()

    # PDF
    if ext == "pdf":
        reader = PdfReader(file)
        text = "".join([page.extract_text() + "\n" for page in reader.pages])
        return text.strip()
    # TXT
    elif ext == "txt":
        if hasattr(file, "read"):
            return file.read().decode("utf-8")
        else:
            with open(file, "r", encoding="utf-8") as f:
                return f.read()
    # DOCX
    elif ext == "docx":
        import docx, io
        if hasattr(file, "read"):
            doc = docx.Document(io.BytesIO(file.read()))
        else:
            doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""